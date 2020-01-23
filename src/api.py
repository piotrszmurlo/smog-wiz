import requests
import time
import csv
from datetime import datetime
import asyncio
import docx
import os
import json

dirname = os.path.dirname(__file__) + '/cache/'
current_date = datetime.now()
current_date_str = current_date.strftime("%Y-%m-%d %H:%M:%S")



def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


class API(object):

    @staticmethod
    async def get_stations() -> list:
        """returns a json list of all available stations; fetches data from GIOS API"""
        try:
            cached_response = docx.Document(dirname + 'get_stations.docx')
            date = cached_response.paragraphs[0].text
            data = cached_response.paragraphs[1].text
            date_object = datetime.strptime(date, '%Y-%m-%d %H:%M:%S')
            time_diff = (current_date - date_object).total_seconds()//60
            if time_diff >= 10: #set the amount of time needed to make cache stale
                response = requests.get('http://api.gios.gov.pl/pjp-api/rest/station/findAll').text
                delete_paragraph(cached_response.paragraphs[1])
                delete_paragraph(cached_response.paragraphs[0])
                cached_response.add_paragraph(current_date_str)
                cached_response.add_paragraph(response)
                cached_response.save(dirname + 'get_stations.docx')
                return json.loads(response)
            else:
                return json.loads(data)
        except docx.opc.exceptions.PackageNotFoundError:
            print(f'(cache) get_stations.docx not found. Creating new file...')
            doc = docx.Document()
            doc.add_paragraph('2000-01-21 16:16:09') #temporary date to match time format to match time format
            doc.add_paragraph('null')
            doc.save(dirname + f'get_stations.docx')
            return await API.get_stations()
        except Exception as e:
            time.sleep(1)
            print(e)
            print(f'...fetch retry...')
            return await API.get_stations
    
    @staticmethod
    async def get_measuring_stands_for_station(station_id:int) -> list:
        """returns a json list of all measuring stands for station specified by id in the argument; fetches data from GIOS API"""
        try:
            cached_response = docx.Document(dirname + f'get_measuring_stands_for_station({station_id}).docx')
            date = cached_response.paragraphs[0].text
            data = cached_response.paragraphs[1].text
            date_object = datetime.strptime(date, '%Y-%m-%d %H:%M:%S')
            time_diff = (current_date - date_object).total_seconds()//60
            if time_diff >=10: #set the amount of time needed to make cache stale
                response = requests.get(f'http://api.gios.gov.pl/pjp-api/rest/station/sensors/{station_id}').text
                delete_paragraph(cached_response.paragraphs[1])
                delete_paragraph(cached_response.paragraphs[0])
                cached_response.add_paragraph(current_date_str)
                cached_response.add_paragraph(response)
                cached_response.save(dirname + f'get_measuring_stands_for_station({station_id}).docx')
                return json.loads(response)
            else:
                return json.loads(data)
        except docx.opc.exceptions.PackageNotFoundError:
            print(f'get_measuring_stands_for_station({station_id}).docx not found. Creating new file...')
            doc = docx.Document()
            doc.add_paragraph('2000-01-21 16:16:09') #temporary date to match time format
            doc.add_paragraph('null')
            doc.save(dirname + f'get_measuring_stands_for_station({station_id}).docx')
            return await API.get_measuring_stands_for_station(station_id)
        except TypeError:
            print(f'Station with such id does not exist')

    @staticmethod
    async def get_measuring_stand_data(stand_id:int) -> list:
        """returns a json list of all data from measuring stand specified by id in the argument; fetches data from GIOS API"""
        try:
            cached_response = docx.Document(dirname + f'get_measuring_stand_data({stand_id}).docx')
            date = cached_response.paragraphs[0].text
            data = cached_response.paragraphs[1].text
            date_object = datetime.strptime(date, '%Y-%m-%d %H:%M:%S')
            time_diff = (current_date - date_object).total_seconds()//60
            if time_diff >=10: #set the amount of time needed to make cache stale
                response = requests.get(f'http://api.gios.gov.pl/pjp-api/rest/data/getData/{stand_id}').text
                delete_paragraph(cached_response.paragraphs[1])
                delete_paragraph(cached_response.paragraphs[0])
                cached_response.add_paragraph(current_date_str)
                cached_response.add_paragraph(response)
                cached_response.save(dirname + f'get_measuring_stand_data({stand_id}).docx')
                return json.loads(response)
            else:
                return json.loads(data)
        except docx.opc.exceptions.PackageNotFoundError:
            print(f'get_measuring_stand_data({stand_id}).docx not found. Creating new file...')
            doc = docx.Document()
            doc.add_paragraph('2000-01-21 16:16:09') #temporary date to match time format to match time formar
            doc.add_paragraph('null')
            doc.save(dirname + f'get_measuring_stand_data({stand_id}).docx')
            return await API.get_measuring_stand_data(stand_id)
        except TypeError:
            print(f'Stand with such id does not exist')
            